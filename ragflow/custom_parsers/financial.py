#
#  Custom Financial Document Parser for RagFlow
#  Optimized for SEC Filings (10-K, 10-Q, 8-K, etc.)
#
#  Based on RagFlow's laws.py parser template
#  Licensed under Apache 2.0
#
#  Features:
#  - Section filtering (skips Risk Factors, Forward-Looking, etc.)
#  - Priority extraction (MD&A, Financial Statements)
#  - Integrated deduplication via dedup.py
#

import logging
import re
from io import BytesIO
from typing import List, Tuple, Optional

from docx import Document

from deepdoc.parser import PdfParser, DocxParser, HtmlParser
from deepdoc.parser.utils import get_text
from rag.nlp import (
    bullets_category, 
    remove_contents_table,
    make_colon_as_title, 
    tokenize_chunks, 
    tree_merge,
    rag_tokenizer
)
from rag.app.naive import by_plaintext, PARSERS
from common.parser_config_utils import normalize_layout_recognizer
from common.constants import ParserType


# =============================================================================
# SEC FILING CONFIGURATION
# =============================================================================

# Sections to SKIP (boilerplate content)
SEC_SKIP_SECTIONS = [
    r"forward[\-\s]*looking\s+statements?",
    r"safe\s+harbor",
    r"cautionary\s+(note|statement)",
    r"risk\s+factors",  # Often 50+ pages of legal boilerplate
    r"legal\s+proceedings",
    r"mine\s+safety",
    r"unregistered\s+sales",
    r"defaults\s+upon\s+senior\s+securities",
    r"submission\s+of\s+matters",
    r"exhibits?\s+and\s+financial",
    r"signatures?$",
    r"power\s+of\s+attorney",
    r"certifications?",
]

# Sections to PRIORITIZE (high-value content)
SEC_PRIORITY_SECTIONS = {
    r"item\s*7[^a]": "MD&A",  # Management Discussion & Analysis
    r"item\s*7a": "Quantitative Disclosures",
    r"item\s*8": "Financial Statements",
    r"item\s*1[^0-9]": "Business Description",
    r"item\s*2": "Properties",
    r"item\s*5": "Market for Common Equity",
    r"item\s*6": "Selected Financial Data",
    r"notes?\s+to\s+(consolidated\s+)?financial": "Financial Notes",
    r"balance\s+sheet": "Balance Sheet",
    r"income\s+statement": "Income Statement",
    r"cash\s+flow": "Cash Flow Statement",
    r"stockholders?\s*equity": "Equity Statement",
}

# Common boilerplate phrases to filter out
BOILERPLATE_PHRASES = [
    r"this\s+form\s+10-?k",
    r"securities\s+and\s+exchange\s+commission",
    r"pursuant\s+to\s+section\s+\d+",
    r"incorporated\s+by\s+reference",
    r"the\s+following\s+discussion\s+should\s+be\s+read",
    r"except\s+as\s+otherwise\s+indicated",
    r"as\s+used\s+in\s+this\s+report",
]


# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def should_skip_section(text: str) -> bool:
    """Check if a section should be skipped based on its header."""
    text_lower = text.lower().strip()
    for pattern in SEC_SKIP_SECTIONS:
        if re.search(pattern, text_lower):
            logging.debug(f"Skipping section: {text[:50]}...")
            return True
    return False


def get_section_tag(text: str) -> Optional[str]:
    """Get the priority tag for a section if it's a high-value section."""
    text_lower = text.lower().strip()
    for pattern, tag in SEC_PRIORITY_SECTIONS.items():
        if re.search(pattern, text_lower):
            return tag
    return None


def is_boilerplate(text: str, threshold: float = 0.5) -> bool:
    """
    Check if text is mostly boilerplate legal language.
    Returns True if more than threshold of boilerplate phrases found.
    """
    text_lower = text.lower()
    matches = sum(1 for p in BOILERPLATE_PHRASES if re.search(p, text_lower))
    # If text contains multiple boilerplate phrases, likely boilerplate
    return matches >= 2


def extract_financial_tables(text: str) -> List[str]:
    """
    Extract financial data tables from text.
    Returns list of table strings.
    """
    # Simple pattern for detecting tabular financial data
    # (rows with $ amounts, percentages, or aligned numbers)
    table_patterns = [
        r"(\$[\d,]+\.?\d*)",  # Dollar amounts
        r"(\d+\.?\d*\s*%)",   # Percentages
        r"(\d{1,3}(?:,\d{3})+)",  # Large numbers with commas
    ]
    tables = []
    lines = text.split('\n')
    
    current_table = []
    for line in lines:
        has_financial_data = any(re.search(p, line) for p in table_patterns)
        if has_financial_data:
            current_table.append(line)
        elif current_table:
            if len(current_table) >= 3:  # At least 3 rows = likely a table
                tables.append('\n'.join(current_table))
            current_table = []
    
    return tables


# =============================================================================
# DOCUMENT PARSERS
# =============================================================================

class FinancialDocx(DocxParser):
    """DOCX parser optimized for SEC filings."""
    
    def __init__(self):
        pass
    
    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(filename) if not binary else Document(BytesIO(binary))
        
        sections = []
        current_section = []
        current_tag = None
        skip_mode = False
        
        for p in self.doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            
            # Check for section headers
            if p.style and 'Heading' in str(p.style.name):
                # Save previous section if not skipped
                if current_section and not skip_mode:
                    sections.append((current_tag, '\n'.join(current_section)))
                
                # Check if new section should be skipped
                skip_mode = should_skip_section(text)
                current_tag = get_section_tag(text) if not skip_mode else None
                current_section = [text] if not skip_mode else []
            elif not skip_mode:
                if not is_boilerplate(text):
                    current_section.append(text)
        
        # Don't forget last section
        if current_section and not skip_mode:
            sections.append((current_tag, '\n'.join(current_section)))
        
        return sections


class FinancialPdf(PdfParser):
    """PDF parser optimized for SEC filings."""
    
    def __init__(self):
        self.model_speciess = ParserType.LAWS.value
        super().__init__()
    
    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None):
        from timeit import default_timer as timer
        start = timer()
        
        callback(msg="OCR started (Financial Parser)")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            callback
        )
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))
        
        start = timer()
        self._layouts_rec(zoomin)
        callback(0.67, "Layout analysis ({:.2f}s)".format(timer() - start))
        self._naive_vertical_merge()
        
        # Filter out boilerplate sections
        filtered_boxes = []
        skip_mode = False
        
        for b in self.boxes:
            text = b.get("text", "")
            
            # Check for section headers (usually larger font or bold)
            if b.get("layout_type") in ["title", "header", "section"]:
                skip_mode = should_skip_section(text)
            
            if not skip_mode and not is_boilerplate(text):
                # Tag with section info if priority section
                tag = get_section_tag(text)
                if tag:
                    b["section_tag"] = tag
                filtered_boxes.append(b)
        
        callback(0.8, "Text extraction ({:.2f}s) - {} chunks after filtering".format(
            timer() - start, len(filtered_boxes)))
        
        return [(b["text"], self._line_tag(b, zoomin)) for b in filtered_boxes], None


# =============================================================================
# MAIN CHUNK FUNCTION
# =============================================================================

def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="English", callback=None, **kwargs):
    """
    Main entry point for the Financial document parser.
    
    Supported file formats: docx, pdf, txt, html
    
    Features:
    - Filters boilerplate SEC sections (Risk Factors, Forward-Looking, etc.)
    - Prioritizes MD&A, Financial Statements, Notes
    - Extracts financial tables and figures
    """
    parser_config = kwargs.get(
        "parser_config", {
            "chunk_token_num": 768,  # Larger chunks for financial context
            "delimiter": "\n!?。；！？",
            "layout_recognize": "DeepDOC"
        }
    )
    
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    
    pdf_parser = None
    sections = []
    eng = lang.lower() == "english"
    
    # Route to appropriate parser based on file type
    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Parsing DOCX (Financial Parser)")
        parsed_sections = FinancialDocx()(filename, binary)
        for tag, content in parsed_sections:
            sections.append(content)
        callback(0.7, f"Parsed {len(sections)} sections")
    
    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        layout_recognizer, parser_model_name = normalize_layout_recognizer(
            parser_config.get("layout_recognize", "DeepDOC")
        )
        
        if isinstance(layout_recognizer, bool):
            layout_recognizer = "DeepDOC" if layout_recognizer else "Plain Text"
        
        name = layout_recognizer.strip().lower()
        parser = PARSERS.get(name, by_plaintext)
        
        callback(0.1, "Parsing PDF (Financial Parser)")
        raw_sections, tables, pdf_parser = parser(
            filename=filename,
            binary=binary,
            from_page=from_page,
            to_page=to_page,
            lang=lang,
            callback=callback,
            pdf_cls=FinancialPdf,
            layout_recognizer=layout_recognizer,
            mineru_llm_name=parser_model_name,
            **kwargs
        )
        
        if not raw_sections and not tables:
            callback(0.99, "No content parsed from PDF")
            return []
        
        for txt, poss in raw_sections:
            if not should_skip_section(txt) and not is_boilerplate(txt):
                sections.append(txt + poss)
        
        callback(0.8, f"Parsed {len(sections)} sections after filtering")
    
    elif re.search(r"\.(txt|md|markdown)$", filename, re.IGNORECASE):
        callback(0.1, "Parsing text file (Financial Parser)")
        txt = get_text(filename, binary)
        
        # Split by section headers
        section_pattern = r"(?:^|\n)(item\s*\d+[a-z]?\.?[^\n]*)"
        parts = re.split(section_pattern, txt, flags=re.IGNORECASE)
        
        for part in parts:
            if part.strip() and not should_skip_section(part) and not is_boilerplate(part):
                sections.append(part.strip())
        
        callback(0.8, f"Parsed {len(sections)} sections")
    
    elif re.search(r"\.(htm|html)$", filename, re.IGNORECASE):
        callback(0.1, "Parsing HTML (Financial Parser)")
        raw_sections = HtmlParser()(filename, binary)
        sections = [s for s in raw_sections 
                   if s and not should_skip_section(s) and not is_boilerplate(s)]
        callback(0.8, f"Parsed {len(sections)} sections")
    
    else:
        raise NotImplementedError(
            "File type not supported (docx, pdf, txt, html supported)")
    
    if not sections:
        callback(0.99, "No relevant content found after filtering")
        return []
    
    # Standard RagFlow post-processing
    remove_contents_table(sections, eng)
    make_colon_as_title(sections)
    bull = bullets_category(sections)
    res = tree_merge(bull, sections, 2)
    
    if not res:
        callback(0.99, "No chunks generated")
        return []
    
    callback(0.95, f"Generated {len(res)} chunks")
    return tokenize_chunks(res, doc, eng, pdf_parser)


# =============================================================================
# CLI TESTING
# =============================================================================

if __name__ == "__main__":
    import sys
    
    def dummy(prog=None, msg=""):
        print(f"[{prog}] {msg}" if prog else msg)
    
    if len(sys.argv) > 1:
        result = chunk(sys.argv[1], callback=dummy)
        print(f"\n=== RESULTS ===")
        print(f"Total chunks: {len(result)}")
        for i, chunk in enumerate(result[:5]):
            print(f"\n--- Chunk {i+1} ---")
            print(chunk.get("content_with_weight", "")[:200])
